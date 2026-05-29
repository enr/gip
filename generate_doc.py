#!/usr/bin/env python3
"""Generate functional documentation for gip in Word format."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TODAY = datetime.date.today().strftime("%d %B %Y")

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles helpers ---
def set_heading_style(para, size, bold=True, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para.alignment = align
    run = para.runs[0] if para.runs else para.add_run()
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_box(doc, label, text, bg_color=None):
    """Add a colored info box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    label_run = p.add_run(f"{label} ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    text_run = p.add_run(text)
    text_run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ==========================================
# 1. COVER PAGE
# ==========================================
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run("gip")
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Gestore Multi-Repository Git per Team di Sviluppo")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()

doc_type = doc.add_paragraph()
doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = doc_type.add_run("Documentazione Funzionale — Destinata al Team Commerciale")
rt.font.size = Pt(13)
rt.bold = True

doc.add_paragraph()

tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rtag = tag_p.add_run("★  Estratta dal codice sorgente — Source of Truth Edition  ★")
rtag.font.size = Pt(11)
rtag.font.color.rgb = RGBColor(0xC0, 0x50, 0x10)
rtag.italic = True

doc.add_paragraph()

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rver = ver_p.add_run(f"Versione 1.0  |  {TODAY}")
rver.font.size = Pt(11)
rver.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ==========================================
# 2. COS'È IL PRODOTTO
# ==========================================
add_h1(doc, "1. Cos'è gip")

add_body(doc, (
    "gip è uno strumento da riga di comando che consente ai team di sviluppo di "
    "gestire in modo centralizzato un insieme di repository Git. "
    "Anziché entrare manualmente in ciascuna cartella di progetto per eseguire "
    "aggiornamenti o verificare lo stato del codice, con gip è possibile eseguire "
    "queste operazioni su tutti i repository registrati con un singolo comando."
))
doc.add_paragraph()

add_body(doc, (
    "La configurazione è semplice: l'utente crea un file di testo che elenca i "
    "progetti (con nome, indirizzo remoto e cartella locale), e da quel momento "
    "può operare su tutti contemporaneamente — ottenendo aggiornamenti, "
    "controllando lo stato delle modifiche, eseguendo comandi personalizzati e "
    "molto altro — con piena visibilità sull'esito di ogni operazione."
))
doc.add_paragraph()

add_body(doc, (
    "gip è progettato per lavorare in parallelo: le operazioni su più repository "
    "vengono eseguite contemporaneamente (fino a 32 in parallelo), riducendo "
    "drasticamente i tempi di attesa anche per team con decine o centinaia di progetti."
))
doc.add_paragraph()

add_box(doc,
    "Flusso principale:",
    "Configura una lista di repository → esegui un comando gip → ottieni "
    "il riepilogo dell'operazione su tutti i progetti in pochi secondi.")

doc.add_page_break()

# ==========================================
# 3. A CHI È RIVOLTO
# ==========================================
add_h1(doc, "2. A Chi È Rivolto")

add_body(doc, "gip si rivolge a professionisti e team che gestiscono quotidianamente più repository Git:")
doc.add_paragraph()

bullets = [
    "Team di sviluppo software con progetti suddivisi in più repository (architettura a microservizi, monorepo parziali, librerie condivise)",
    "Sviluppatori individuali che mantengono contemporaneamente più progetti open-source o commerciali",
    "Team di DevOps e ingegneria della piattaforma che devono mantenere sincronizzati ambienti di sviluppo complessi",
    "Agenzie di sviluppo che gestiscono codice per clienti multipli, ciascuno con il proprio repository",
    "Aziende che distribuiscono un prodotto modulare composto da più componenti Git indipendenti",
]
for b in bullets:
    add_bullet(doc, b)

doc.add_paragraph()
add_body(doc, "Settori di riferimento principali:")
doc.add_paragraph()

sectors = [
    "Sviluppo software prodotto (ISV — Independent Software Vendor)",
    "Servizi IT e system integration",
    "Fintech, startup tecnologiche, aziende con cultura engineering-first",
    "Enti pubblici e grandi organizzazioni con più team di sviluppo",
]
for s in sectors:
    add_bullet(doc, s)

doc.add_page_break()

# ==========================================
# 4. VANTAGGI / DIFFERENZIATORI
# ==========================================
add_h1(doc, "3. Vantaggi e Differenziatori")

add_table(doc,
    ["Vantaggio", "Descrizione"],
    [
        ["Visione unificata", "Un solo comando mostra lo stato di tutti i repository: chi ha modifiche non salvate, chi è indietro rispetto al server remoto, su quale ramo si trova ciascun progetto."],
        ["Risparmio di tempo", "Le operazioni vengono eseguite in parallelo su tutti i repository. Aggiornare 20 progetti richiede lo stesso tempo di aggiornarne uno solo."],
        ["Nessuna curva di apprendimento", "Basato su Git standard. Non introduce nuovi concetti né sostituisce il flusso di lavoro esistente — si aggiunge sopra."],
        ["Automazione flessibile", "Il comando 'exec' permette di eseguire qualsiasi script o programma in ogni cartella di progetto, aprendo la porta all'automazione completa del ciclo di sviluppo."],
        ["Output machine-readable", "Modalità JSON integrata per integrare gip in pipeline CI/CD, dashboard aziendali o strumenti di monitoring personalizzati."],
        ["Controllo totale", "Modalità dry-run per vedere cosa succederebbe prima di eseguire: nessuna sorpresa, nessuna modifica accidentale."],
        ["Multipiattaforma", "Funziona su Linux, macOS e Windows senza dipendenze esterne oltre a Git stesso."],
        ["Configurazione portatile", "Un singolo file di testo (YAML o JSON) descrive l'intero ambiente di sviluppo. Si condivide nel team, si mette in versione, si porta su qualsiasi macchina."],
    ]
)

doc.add_page_break()

# ==========================================
# 5. CHI USA IL SISTEMA
# ==========================================
add_h1(doc, "4. Chi Usa il Sistema")

add_table(doc,
    ["Ruolo", "Chi è nella realtà aziendale", "Cosa può fare con gip"],
    [
        ["Sviluppatore", "Ingegnere software del team di prodotto", "Aggiorna tutti i repository con un comando, verifica le modifiche locali, vede su quale ramo lavora ogni progetto, esegue comandi personalizzati in tutti i progetti."],
        ["Tech Lead / Responsabile Tecnico", "Senior engineer o responsabile di un gruppo di sviluppo", "Monitora lo stato di sincronizzazione dell'intero parco repository, identifica chi è rimasto indietro rispetto al server, verifica l'allineamento dei rami."],
        ["DevOps / Platform Engineer", "Ingegnere responsabile dell'infrastruttura e dei processi", "Integra gip in pipeline automatizzate, utilizza l'output JSON per alimentare dashboard di monitoraggio, automatizza operazioni di manutenzione su tutti i progetti."],
        ["Sviluppatore Freelance / Consulente", "Professionista che gestisce più clienti o progetti in parallelo", "Organizza i repository per cliente usando i tag, aggiorna selettivamente solo i progetti rilevanti, mantiene l'ambiente di lavoro sempre sincronizzato."],
    ]
)

doc.add_page_break()

# ==========================================
# 6. FUNZIONALITÀ PRINCIPALI
# ==========================================
add_h1(doc, "5. Funzionalità Principali")

# 6.1 Configurazione e censimento
add_h2(doc, "5.1 Configurazione dell'Ambiente di Lavoro")

add_body(doc, (
    "Prima di utilizzare gip, l'utente definisce quali repository gestire. "
    "Questa configurazione può essere creata manualmente oppure generata "
    "automaticamente: gip è in grado di analizzare una cartella del computer, "
    "individuare tutti i repository Git presenti e produrre automaticamente il file "
    "di configurazione con tutte le informazioni necessarie."
))
doc.add_paragraph()

add_body(doc, "Il processo di configurazione automatica funziona così:")
add_numbered(doc, "L'utente indica a gip una cartella del proprio computer da analizzare")
add_numbered(doc, "gip percorre tutte le sottocartelle alla ricerca di repository Git")
add_numbered(doc, "Per ogni repository trovato, recupera automaticamente l'indirizzo del server remoto")
add_numbered(doc, "Genera il file di configurazione con tutti i repository censiti")
add_numbered(doc, "L'utente può personalizzare il file risultante (aggiungere etichette, modificare policy, disabilitare progetti)")
doc.add_paragraph()

add_box(doc,
    "Valore per il cliente:",
    "Censire un ambiente di sviluppo con 50 repository richiede pochi secondi invece "
    "di ore di lavoro manuale. Il file generato diventa la fonte unica di verità "
    "per l'intero team.")

doc.add_paragraph()

# 6.2 Aggiornamento dei repository
add_h2(doc, "5.2 Aggiornamento Massivo dei Repository")

add_body(doc, (
    "Il comando di aggiornamento scarica le ultime modifiche dal server remoto "
    "per tutti i repository configurati, eseguendo le operazioni in parallelo. "
    "Il comportamento di aggiornamento per ciascun progetto è configurabile "
    "tramite una 'politica di aggiornamento'."
))
doc.add_paragraph()

add_body(doc, "Sono disponibili tre politiche:")

add_table(doc,
    ["Politica", "Comportamento"],
    [
        ["Predefinita", "Aggiorna il repository se la cartella locale è già presente. Se la cartella non esiste, segnala il problema senza agire."],
        ["Sempre (always)", "Aggiorna il repository. Se la cartella locale non esiste ancora, scarica automaticamente l'intero repository dal server (clonazione)."],
        ["Mai (never)", "Il repository viene escluso da qualsiasi operazione di aggiornamento o scaricamento remoto. Utile per progetti in cui non si vuole mai sovrascrivere il codice locale."],
    ]
)

add_body(doc, (
    "In caso di cartella mancante, gip avvisa l'utente. Con il parametro 'tutti' "
    "attivo, i repository mancanti vengono automaticamente scaricati dal server "
    "remoto, rendendo possibile ricostruire un intero ambiente di sviluppo su una "
    "nuova macchina con un solo comando."
))
doc.add_paragraph()

add_box(doc,
    "Valore per il cliente:",
    "Un nuovo sviluppatore che si unisce al team può avere tutti i repository "
    "dell'organizzazione sulla propria macchina in pochi minuti, senza dover "
    "copiare decine di indirizzi manualmente.")

doc.add_paragraph()

# 6.3 Verifica stato modifiche
add_h2(doc, "5.3 Verifica dello Stato delle Modifiche")

add_body(doc, (
    "gip permette di avere in un colpo d'occhio la fotografia dello stato di tutti "
    "i repository: quali file sono stati modificati, quali modifiche sono pronte "
    "per essere salvate, quali file nuovi non sono ancora tracciati."
))
doc.add_paragraph()

add_body(doc, "Sono disponibili due livelli di dettaglio:")

add_table(doc,
    ["Modalità", "Cosa mostra"],
    [
        ["Stato standard", "Mostra solo i file modificati rispetto all'ultima versione salvata. Ideale per una rapida verifica quotidiana."],
        ["Stato completo", "Mostra sia i file modificati sia quelli nuovi non ancora presi in carico. Utile prima di un'operazione di aggiornamento dal server."],
    ]
)

add_body(doc, (
    "Con la modalità estesa, gip mostra anche indicatori compatti per ogni repository: "
    "se ha modifiche pronte da salvare (+), modifiche non ancora preparate (*), "
    "file nuovi non tracciati (?), o modifiche messe temporaneamente da parte ($). "
    "Questo consente di filtrare e lavorare solo sui repository che richiedono attenzione."
))
doc.add_paragraph()

add_box(doc,
    "Valore per il cliente:",
    "Prima di una release o di una revisione del codice, il responsabile tecnico "
    "può identificare in secondi quali repository hanno lavoro in corso non completato, "
    "evitando rilasci accidentali di codice non finalizzato.")

doc.add_paragraph()

# 6.4 Sincronizzazione remota (fetch)
add_h2(doc, "5.4 Scaricamento delle Informazioni Remote")

add_body(doc, (
    "Il comando di scaricamento aggiorna le informazioni sui repository remoti "
    "senza modificare il codice locale. Questo permette a gip di calcolare "
    "con precisione quante modifiche ogni repository ha ricevuto sul server "
    "rispetto alla versione locale, senza rischiare di sovrascrivere il lavoro in corso."
))
doc.add_paragraph()

add_body(doc, (
    "Questa operazione è il prerequisito consigliato prima di usare il filtro "
    "'in ritardo' (vedi sezione 5.6): aggiornare prima le informazioni remote "
    "garantisce che il confronto sia accurato."
))
doc.add_paragraph()

add_box(doc,
    "Valore per il cliente:",
    "I team con rilasci frequenti possono programmare questa operazione in modo "
    "automatico per avere sempre informazioni aggiornate sulla divergenza tra "
    "codice locale e server, senza interrompere il lavoro in corso.")

doc.add_paragraph()

# 6.5 Panoramica dei rami attivi
add_h2(doc, "5.5 Panoramica dei Rami di Sviluppo")

add_body(doc, (
    "Il comando di visualizzazione dei rami mostra, per ogni repository, "
    "su quale ramo di sviluppo si trova attualmente il codice locale. "
    "In modalità estesa, la visualizzazione include:"
))

add_bullet(doc, "Il ramo attivo di ciascun repository")
add_bullet(doc, "Lo stato di sincronizzazione rispetto al server (in anticipo, in ritardo, allineato, divergente)")
add_bullet(doc, "Gli indicatori di modifiche non salvate")
add_bullet(doc, "Il messaggio e la data dell'ultima modifica salvata")
doc.add_paragraph()

add_body(doc, (
    "Lo stato di sincronizzazione viene rappresentato con simboli intuitivi: "
    "'allineato' quando il codice locale e il server coincidono, "
    "freccia su con numero quando il locale è avanti rispetto al server, "
    "freccia giù con numero quando il server ha aggiornamenti non ancora scaricati, "
    "entrambe le frecce quando i due ambienti sono divergenti."
))
doc.add_paragraph()

add_box(doc,
    "Valore per il cliente:",
    "In un team che lavora su feature branch, il tech lead può verificare "
    "istantaneamente su quale ramo lavora ogni sviluppatore e se ci sono "
    "repository rimasti su rami di sviluppo che avrebbero dovuto essere chiusi.")

doc.add_paragraph()

# 6.6 Filtri avanzati
add_h2(doc, "5.6 Filtri Avanzati per Operazioni Mirate")

add_body(doc, (
    "Ogni comando di gip supporta filtri che permettono di operare solo su un "
    "sottoinsieme dei repository, aumentando la precisione e riducendo il tempo "
    "di elaborazione."
))
doc.add_paragraph()

add_table(doc,
    ["Filtro", "Effetto"],
    [
        ["Per etichetta (tag)", "Esegue l'operazione solo sui repository che portano almeno una delle etichette specificate. Più etichette si combinano con logica 'almeno una' (OR). Esempio: --tag lavoro,frontend agisce su tutti i repository con etichetta 'lavoro' oppure 'frontend'."],
        ["Solo con modifiche (dirty)", "Esegue l'operazione solo sui repository che hanno almeno una modifica locale non ancora salvata in modo definitivo."],
        ["Solo in ritardo (behind)", "Esegue l'operazione solo sui repository che hanno aggiornamenti sul server remoto non ancora scaricati localmente."],
    ]
)

add_body(doc, (
    "I filtri possono essere combinati e si applicano a tutti i comandi principali: "
    "aggiornamento, scaricamento, visualizzazione rami, visualizzazione stato e "
    "esecuzione comandi personalizzati."
))
doc.add_paragraph()

add_box(doc,
    "Valore per il cliente:",
    "Un team che divide i repository per cliente o area di prodotto può, "
    "con un singolo comando, aggiornare tutti e soli i repository del cliente "
    "X senza toccare quelli degli altri clienti.")

doc.add_paragraph()

# 6.7 Esecuzione comandi personalizzati
add_h2(doc, "5.7 Esecuzione di Comandi Personalizzati")

add_body(doc, (
    "La funzionalità più potente di gip è la possibilità di eseguire qualsiasi "
    "comando del sistema operativo o script in ogni cartella di repository, "
    "in modo parallelo, raccogliendo e mostrando i risultati in modo uniforme."
))
doc.add_paragraph()

add_body(doc, "Esempi di utilizzo pratico:")
add_bullet(doc, "Eseguire i test automatizzati di tutti i progetti con un solo comando")
add_bullet(doc, "Aggiornare le dipendenze esterne in tutti i repository")
add_bullet(doc, "Cercare occorrenze di un termine specifico nel codice di tutti i progetti")
add_bullet(doc, "Avviare processi di compilazione o packaging in parallelo")
add_bullet(doc, "Eseguire script di migrazione o manutenzione su tutto il parco progetti")
doc.add_paragraph()

add_box(doc,
    "Valore per il cliente:",
    "Un team che deve applicare una modifica di sicurezza a tutti i propri "
    "progetti può automatizzare l'operazione e completarla in minuti anziché "
    "in giorni di lavoro manuale ripetitivo.")

doc.add_paragraph()

# 6.8 Modalità di output
add_h2(doc, "5.8 Modalità di Visualizzazione e Integrazione")

add_body(doc, (
    "gip offre diverse modalità di output per adattarsi a contesti diversi, "
    "dal terminale del singolo sviluppatore ai sistemi automatizzati."
))
doc.add_paragraph()

add_table(doc,
    ["Modalità", "Quando usarla"],
    [
        ["Normale (default)", "Utilizzo quotidiano: output colorato con barra di avanzamento in tempo reale, riepilogo finale con conteggio successi/errori/saltati e durata totale."],
        ["Silenziosa (--quiet)", "Quando si vuole solo il riepilogo finale senza dettaglio progetto per progetto. Utile per operazioni su molti repository."],
        ["JSON (--json)", "Integrazione con sistemi automatizzati, pipeline CI/CD, dashboard. Produce un oggetto JSON strutturato con tutti i dati dell'operazione."],
        ["Simulazione (--noop)", "Prima di eseguire operazioni impattanti: mostra esattamente cosa verrebbe fatto senza modificare nulla. Ideale per verificare configurazioni prima del deploy."],
    ]
)

add_box(doc,
    "Valore per il cliente:",
    "La modalità JSON permette di costruire dashboard aziendali che mostrano "
    "in tempo reale lo stato di sincronizzazione di tutti i repository, "
    "senza dover accedere ai singoli strumenti Git.")

doc.add_page_break()

# ==========================================
# 7. SCENARI D'USO COMMERCIALI
# ==========================================
add_h1(doc, "6. Scenari d'Uso Commerciali")

add_h2(doc, "Scenario 1 — Agenzia di Sviluppo con Portfolio Multi-Cliente")

add_body(doc, (
    "Un'agenzia con 12 sviluppatori gestisce 40 repository distribuiti tra 8 clienti. "
    "Ogni mattina gli sviluppatori devono scaricare le ultime modifiche prima di "
    "iniziare il lavoro. Con gip, il responsabile tecnico ha configurato i repository "
    "di ogni cliente con etichette dedicate. Ogni sviluppatore aggiorna il proprio "
    "set di lavoro con un singolo comando, filtrando per etichetta cliente. "
    "L'operazione che prima richiedeva 15-20 minuti di apertura terminali e comandi "
    "manuali ora richiede meno di 2 minuti."
))
doc.add_paragraph()

add_box(doc, "Beneficio principale:",
    "Risparmio di 1 ora/giorno per sviluppatore su operazioni di routine, "
    "con eliminazione completa degli errori da 'ho aggiornato il repository sbagliato'.")

doc.add_paragraph()

add_h2(doc, "Scenario 2 — Team di Piattaforma in Azienda Fintech")

add_body(doc, (
    "Un'azienda fintech con 6 microservizi in produzione riceve una notifica di "
    "vulnerabilità di sicurezza in una libreria condivisa. Il team di piattaforma "
    "deve aggiornare la dipendenza in tutti i servizi. Usando gip con il comando "
    "di esecuzione personalizzata, l'ingegnere esegue il comando di aggiornamento "
    "dipendenze su tutti e 6 i repository in parallelo, verifica che i test passino "
    "su tutti, e prepara il rilascio. L'operazione che richiederebbe 2 ore di lavoro "
    "sequenziale viene completata in 20 minuti."
))
doc.add_paragraph()

add_box(doc, "Beneficio principale:",
    "Risposta agli incidenti di sicurezza 6 volte più rapida, "
    "con garanzia di uniformità tra tutti i servizi.")

doc.add_paragraph()

add_h2(doc, "Scenario 3 — Sviluppatore Freelance con Clienti Multipli")

add_body(doc, (
    "Un consulente gestisce 15 progetti per 5 clienti diversi. "
    "Ogni progetto ha il proprio repository su piattaforme diverse "
    "(GitHub, GitLab, server privato). Con gip configura una lista unica "
    "con etichette per cliente. Prima di una riunione con il Cliente A, "
    "aggiorna in 30 secondi tutti e 3 i repository di quel cliente, "
    "verifica quali file ha modificato nell'ultima sessione di lavoro, "
    "e si presenta alla riunione con la situazione del codice perfettamente chiara."
))
doc.add_paragraph()

add_box(doc, "Beneficio principale:",
    "Passaggio istantaneo tra contesti di lavoro diversi, "
    "con visibilità completa sullo stato di ogni progetto.")

doc.add_paragraph()

add_h2(doc, "Scenario 4 — Onboarding Nuovo Sviluppatore")

add_body(doc, (
    "Un nuovo sviluppatore si unisce a un team che gestisce 25 repository. "
    "Tradizionalmente il processo di setup richiedeva un giorno intero: "
    "trovare tutti gli indirizzi dei repository, clonare uno per uno, "
    "configurare le cartelle locali. Con gip il file di configurazione "
    "è già nel repository interno del team. Il nuovo sviluppatore scarica "
    "quel file, esegue un solo comando con il parametro 'tutti' e in "
    "10-15 minuti ha tutti e 25 i repository sul proprio computer, "
    "pronti per lo sviluppo."
))
doc.add_paragraph()

add_box(doc, "Beneficio principale:",
    "Riduzione del tempo di onboarding tecnico da 1 giorno a 15 minuti, "
    "con eliminazione degli errori di configurazione manuale.")

doc.add_page_break()

# ==========================================
# 8. DOMANDE FREQUENTI
# ==========================================
add_h1(doc, "7. Domande Frequenti")

faq = [
    (
        "gip sostituisce Git o lo strumento di versioning che usiamo già?",
        "No, gip non sostituisce Git. Si integra sopra di esso come un livello di "
        "orchestrazione. Gli sviluppatori continuano a usare Git esattamente come prima "
        "per le operazioni sul singolo repository; gip aggiunge la capacità di eseguire "
        "quelle stesse operazioni su tutti i repository con un singolo comando."
    ),
    (
        "Funziona con tutti i provider Git (GitHub, GitLab, Bitbucket, server privato)?",
        "Sì. gip è completamente indipendente dal provider: funziona con qualsiasi "
        "repository Git accessibile dalla rete, sia su piattaforme cloud (GitHub, GitLab, "
        "Bitbucket, Azure DevOps) che su server privati aziendali. La configurazione "
        "usa semplicemente l'indirizzo del repository."
    ),
    (
        "È possibile gestire repository con credenziali diverse?",
        "Sì. gip usa il sistema di autenticazione di Git già configurato sulla macchina "
        "(SSH keys, credential manager, token). Ogni repository usa le proprie credenziali "
        "in modo trasparente, senza necessità di configurazione aggiuntiva in gip."
    ),
    (
        "Si può usare in ambienti CI/CD o pipeline automatizzate?",
        "Sì, è uno degli utilizzi previsti. La modalità JSON produce output strutturato "
        "che può essere consumato da qualsiasi script o sistema di monitoring. "
        "gip è disponibile come eseguibile standalone senza dipendenze aggiuntive, "
        "facilmente installabile in ambienti di automazione."
    ),
    (
        "Cosa succede se un'operazione fallisce su uno dei repository?",
        "gip continua l'esecuzione sugli altri repository e raccoglie tutti gli errori. "
        "Al termine mostra un riepilogo con il conteggio di successi, errori e operazioni "
        "saltate, con la possibilità di visualizzare i dettagli degli errori in una sezione "
        "separata. Il codice di uscita del processo segnala se ci sono stati errori, "
        "permettendo a sistemi automatizzati di reagire di conseguenza."
    ),
    (
        "È possibile fare una prova prima di eseguire operazioni impattanti?",
        "Sì. Il parametro '--noop' attiva una modalità di simulazione completa: gip mostra "
        "esattamente quale comando verrebbe eseguito su ogni repository, senza eseguire "
        "effettivamente nulla. Utile per verificare la configurazione o testare comandi "
        "personalizzati prima di applicarli."
    ),
    (
        "Come si gestisce un repository che non deve mai essere aggiornato automaticamente?",
        "È sufficiente impostare la politica 'never' per quel repository nel file di "
        "configurazione. Da quel momento, qualsiasi operazione di aggiornamento o "
        "scaricamento remoto lo salterà automaticamente, segnalando esplicitamente "
        "che il repository è escluso per policy."
    ),
]

for q, a in faq:
    p = doc.add_paragraph()
    run_q = p.add_run(f"D: {q}")
    run_q.bold = True
    run_q.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    run_q.font.size = Pt(11)

    p2 = doc.add_paragraph()
    run_a = p2.add_run(f"R: {a}")
    run_a.font.size = Pt(11)
    doc.add_paragraph()

doc.add_page_break()

# ==========================================
# 9. NOTA METODOLOGICA
# ==========================================
add_h1(doc, "8. Nota Metodologica")

add_body(doc, (
    "Questo documento è stato generato esclusivamente dall'analisi del codice sorgente "
    "del progetto gip, senza utilizzare specifiche tecniche, file README, documentazione "
    "interna o qualsiasi altro documento preesistente nel repository."
))
doc.add_paragraph()

add_body(doc, (
    "Ogni funzionalità descritta corrisponde a comportamento verificato e implementato "
    "nel codice. Qualsiasi discrepanza con documentazione preesistente va risolta "
    "a favore del codice, che rappresenta la source-of-truth dell'implementazione attuale."
))
doc.add_paragraph()

add_body(doc, (
    "La terminologia utilizzata ha privilegiato il linguaggio operativo e di business "
    "rispetto a quello tecnico, con l'obiettivo di rendere il documento comprensibile "
    "e utile per team commerciali, di prevendita e di supporto."
))
doc.add_paragraph()

# ==========================================
# FOOTER (last paragraph)
# ==========================================
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer_p.add_run(f"gip  |  Documentazione Funzionale v1.0  |  {TODAY}")
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# --- Save ---
output_path = "/home/user/gip/gip_NoSpec.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
